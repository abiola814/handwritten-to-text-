
from django.shortcuts import render, redirect
from django.core.files.storage import FileSystemStorage
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse

import json
import os
import sys
import requests
import time
from io import BytesIO
from main.models import Register,Dart,GetImage
from docx import Document
import imp


import os
import io
from msrest.authentication import CognitiveServicesCredentials
from azure.cognitiveservices.vision.computervision import ComputerVisionClient
from azure.cognitiveservices.vision.computervision.models import OperationStatusCodes,VisualFeatureTypes
import requests
import time
from PIL import Image,ImageDraw,ImageFont


from binascii import a2b_base64
from base64 import b64decode

# missing_env = False
# # Add your Computer Vision subscription key and endpoint to your environment variables.
API_KEY='15a919c8296b48498e7938332a799fe2'
ENDPOINT='https://handwritten22.cognitiveservices.azure.com'
cv_client= ComputerVisionClient(ENDPOINT,CognitiveServicesCredentials(API_KEY))
doc = Document()

def home(request):
#     if request.method == 'POST' and request.FILES['upload']:
#         upload = request.FILES['upload']
#         fss = FileSystemStorage()
#         file = fss.save(upload.name, upload)
#         file_url = fss.url(file)
#         endpoint = 'https://handwritten-to-text2.cognitiveservices.azure.com'
#         subscription_key = '113fad96017f4f6891354e104c42daad'
#         text_recognition_url = endpoint + "/vision/v3.1/read/analyze"

# # Set image_url to the URL of an image that you want to recognize.
#         image_url = 'https://handtotext2.herokuapp.com/media/'+ upload.name
#         headers = {'Ocp-Apim-Subscription-Key': subscription_key}

#         data = {'url': image_url}
#         response = requests.post(
#             text_recognition_url, headers=headers, json=data)
#         response.raise_for_status()

#         # Extracting text requires two API calls: One call to submit the
#         # image for processing, the other to retrieve the text found in the image.

#         # Holds the URI used to retrieve the recognized text.
#         operation_url = response.headers["Operation-Location"]

#         # The recognized text isn't immediately available, so poll to wait for completion.
#         analysis = {}
#         poll = True
#         while (poll):
#             response_final = requests.get(
#                 response.headers["Operation-Location"], headers=headers)
#             analysis = response_final.json()
            
#             print(json.dumps(analysis, indent=4))

#             time.sleep(1)
#             if ("analyzeResult" in analysis):
#                 poll = False
#             if ("status" in analysis and analysis['status'] == 'failed'):
#                 poll = False

#         polygons = []
#         if ("analyzeResult" in analysis):
#             # Extract the recognized text, with bounding boxes.
#             polygons = [(line["boundingBox"], line["text"])
#                         for line in analysis["analyzeResult"]["readResults"][0]["lines"]]

#         # Display the image and overlay it with the extracted text.
    
#         final=''
#         for polygon in polygons:
#             vertices = [(polygon[0][i], polygon[0][i+1])
#                         for i in range(0, len(polygon[0]), 2)]
#             text = polygon[1]
#             wordss=str(text)

#             final=final+ wordss
#             hs=open("media/specs/hesjtigs.txt","a")
#             hs.write(text + "\n")
#             hs.close()
#         files='/media/'+ upload.name

#         g=Register.objects.create(text="specs/hesjtigs.txt")
#         h=Register.objects.get(text="specs/hesjtigs.txt")
#         print(h.text.url)
#         with open("media/specs/hesjt.txt", 'r', encoding='utf-8') as openfile:
#             line = openfile.read()
#             doc.add_paragraph(line)
#             doc.save('file' + ".docx")
#         return render(request, 'base.html', {'file_url':files,'urlu':h.text.url})

    return render(request,"base.html",)
@csrf_exempt
def verify(request):
	info= request.POST.get('word')
	k=Dart.objects.create(text=info)
	p= k.id
	return JsonResponse({'result':p})

def homes(request):
	return render(request,"index.html",)

def image(request):
	return render(request,"image.html",)

@csrf_exempt
def verifyimage(request):
	info= request.FILES.get('image_file')
	pop=Register(text=info)
	pop.save()
	hoh='media/'+ str(pop.text)
	print(hoh)
	response = cv_client.read_in_stream(open(hoh,'rb'),language='en',raw=True)
	time.sleep(5)
	operate= response.headers['Operation-Location']
	operatid=operate.split('/')[-1]
	result= cv_client.get_read_result(operatid)
	print(result)
	print(result.status)
	final=''
	if result.status == OperationStatusCodes.succeeded:
		read_results= result.analyze_result.read_results
		if os.path.exists("media/specs/results.txt"):	
			os.remove("media/specs/results.txt")
		for analyzed_result in read_results:
			for line in analyzed_result.lines:
				print(line.text)
				wordss=str(line.text)

				final=final+ wordss
				print(final)
				hs=open("media/specs/results.txt","a")
				hs.write(line.text + "\n")
				hs.close()
	Register.objects.filter(text="specs/results.txt").delete()
	g=Register.objects.create(text="specs/results.txt")
	h=Register.objects.get(text="specs/results.txt")
	print(h.text.url)
	if os.path.exists('media/specs/result.docx'):
		os.remove("media/specs/result.docx")
	Register.objects.filter(text="specs/result.docx").delete()
	with open("media/specs/results.txt", 'r', encoding='utf-8') as openfile:
		line = openfile.read()
		doc.add_paragraph(line)
		doc.save('media/specs/result' + ".docx")
	dics=Register.objects.create(text="specs/result.docx")
	dicsa=Register.objects.get(text="specs/result.docx")

	return JsonResponse({'result':'info'})
@csrf_exempt
def ver(request,key):
	ki= Dart.objects.get(id=key)
	k=ki.text
	return render(request,'verify.html',{'data':k})

@csrf_exempt
def result(request,key):
	keys = request.POST.get('key')
	ki= Dart.objects.get(id=key)
	k=ki.text	
	kid=str(k)
	ddd=kid
	header, encoded = ddd.split(",", 1)
	data = b64decode(encoded)

	with open("media/specs/image.jpeg", "wb") as f:
		f.write(data)
	g='media/specs/image.jpeg'
	response = cv_client.read_in_stream(open(g,'rb'),language='en',raw=True)
	time.sleep(5)
	operate= response.headers['Operation-Location']
	operatid=operate.split('/')[-1]
	result= cv_client.get_read_result(operatid)
	print(result)
	print(result.status)
	final=''
	if result.status == OperationStatusCodes.succeeded:
		read_results= result.analyze_result.read_results
		if os.path.exists("media/specs/read.txt"):	
			os.remove("media/specs/read.txt")
		for analyzed_result in read_results:
			for line in analyzed_result.lines:
				print(line.text)
				wordss=str(line.text)

				final=final+ wordss
				print(final)
				hs=open("media/specs/read.txt","a")
				hs.write(line.text + "\n")
				hs.close()
	Register.objects.filter(text="specs/read.txt").delete()
	g=Register.objects.create(text="specs/read.txt")
	h=Register.objects.get(text="specs/read.txt")
	print(h.text.url)
	if os.path.exists('media/specs/result.docx'):
		os.remove("media/specs/result.docx")
	Register.objects.filter(text="specs/result.docx").delete()
	with open("media/specs/read.txt", 'r', encoding='utf-8') as openfile:
		line = openfile.read()
		doc.add_paragraph(line)
		doc.save('media/specs/result' + ".docx")
	dics=Register.objects.create(text="specs/result.docx")
	dicsa=Register.objects.get(text="specs/result.docx")

	return JsonResponse({'data':h.text.url,'result':dicsa.text.url})
